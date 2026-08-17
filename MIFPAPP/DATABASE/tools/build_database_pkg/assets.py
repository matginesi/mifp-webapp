#!/usr/bin/env python3
"""Asset download and storage functions."""

import concurrent.futures
import hashlib
import json
import logging
import os
import re
import unicodedata
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse
from tqdm import tqdm

log = logging.getLogger(__name__)

from pypdf import PdfReader

from .config import WEBAPP, DEFAULT_JSONL_DIR, FETCH_UA, ASSETS_DIR


class PermanentAssetDownloadError(Exception):
    """Raised when retrying cannot turn the response into a usable asset."""


def _safe_stem(name):
    if not name:
        return ''
    name = unicodedata.normalize('NFKD', str(name))
    name = name.encode('ascii', 'ignore').decode()
    name = re.sub(r'[^\w\-_.]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name


def _ext_from_url(url):
    """Safe extension extraction from a URL, defaulting to 'bin'."""
    path = urlparse(url).path.rstrip('/')
    if '.' in path:
        ext = path.rsplit('.', 1)[-1]
        if ext and '/' not in ext and len(ext) < 10:
            return ext
    return 'bin'


def _sha256_bytes(data):
    return hashlib.sha256(data).hexdigest()


def _extract_title_from_document(data, extension):
    """Extract title from document content using appropriate library."""
    try:
        # Handle PDF files
        if extension in ['.pdf', 'pdf']:
            if not data.startswith(b'%PDF'):
                return ''
            reader = PdfReader(BytesIO(data))
            if reader.metadata and hasattr(reader.metadata, 'title') and reader.metadata.title:
                return reader.metadata.title
            if len(reader.pages) > 0:
                first_page = reader.pages[0]
                page_text = first_page.extract_text()
                if page_text:
                    for line in page_text.split('\n'):
                        line = line.strip()
                        if line and len(line) > 5 and len(line) < 100:
                            return line[:50].strip()
        
        elif extension in ['.docx', 'docx']:
            if not data.startswith(b'PK'):
                return ''
            from docx import Document
            doc = Document(BytesIO(data))
            
            # Try to get title from document properties
            if doc.core_properties.title:
                return doc.core_properties.title
            
            # Try to get title from first paragraph
            if doc.paragraphs:
                first_para = doc.paragraphs[0]
                title = first_para.text.strip()
                if title and len(title) > 5 and len(title) < 100:
                    return title[:50].strip()
            
            # Try to get title from header
            for section in doc.sections:
                for header in [section.header, section.first_page_header]:
                    if header and header.paragraphs:
                        for para in header.paragraphs:
                            text = para.text.strip()
                            if text and len(text) > 5 and len(text) < 100:
                                return text[:50].strip()
        
        # Handle HTML files
        elif extension in ['.html', 'htm', 'html']:
            import re
            title_match = re.search(rb'<title[^>]*>([^<]+)</title>', data, re.IGNORECASE | re.DOTALL)
            if title_match:
                return title_match.group(1).decode('utf-8', errors='ignore').strip()
    except Exception:
        pass
    
    return ''


def _guess_mime_from_ext(path):
    if path.endswith('.png'):
        return 'image/png'
    elif path.endswith('.jpg') or path.endswith('.jpeg'):
        return 'image/jpeg'
    elif path.endswith('.gif'):
        return 'image/gif'
    elif path.endswith('.svg'):
        return 'image/svg+xml'
    elif path.endswith('.webp'):
        return 'image/webp'
    elif path.endswith('.avif'):
        return 'image/avif'
    elif path.endswith('.ico'):
        return 'image/x-icon'
    elif path.endswith('.tiff') or path.endswith('.tif'):
        return 'image/tiff'
    elif path.endswith('.bmp'):
        return 'image/bmp'
    elif path.endswith('.pdf'):
        return 'application/pdf'
    elif path.endswith('.mp4'):
        return 'video/mp4'
    elif path.endswith('.webm'):
        return 'video/webm'
    elif path.endswith('.mp3'):
        return 'audio/mpeg'
    return 'application/octet-stream'


def _guess_kind_from_ext(ext):
    """Guess the asset kind from file extension."""
    # Remove leading dot if present, convert to lowercase
    ext = ext.lstrip('.')
    ext = ext.lower()
    
    # Map extensions to asset kinds (without dots)
    if ext in {'png', 'jpeg', 'jpg', 'gif', 'svg', 'webp', 'avif', 'ico', 'tiff', 'tif', 'bmp'}:
        return 'image'
    elif ext == 'pdf':
        return 'pdf'
    elif ext in {'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx', 'csv', 'txt'}:
        return 'document'
    elif ext in {'mp4', 'mov'}:
        return 'video'
    elif ext in {'mp3', 'wav'}:
        return 'audio'
    elif ext in {'zip', 'rar', 'tar', 'gz'}:
        return 'archive'
    else:
        return 'other'


def _looks_like_html(data):
    head = (data or b'')[:512].lstrip().lower()
    return head.startswith((b'<!doctype html', b'<html', b'<?xml')) or b'<html' in head[:256]


def _asset_invalid_reason(data, url):
    if not data:
        return "empty response"
    if _looks_like_html(data):
        return "HTML response instead of asset"
    ext = _ext_from_url(url).lower().lstrip('.')
    if ext in {'jpg', 'jpeg'} and not data.startswith(b'\xff\xd8\xff'):
        return "JPEG signature mismatch"
    if ext == 'png' and not data.startswith(b'\x89PNG\r\n\x1a\n'):
        return "PNG signature mismatch"
    if ext == 'gif' and not data.startswith((b'GIF87a', b'GIF89a')):
        return "GIF signature mismatch"
    if ext == 'webp' and not (data.startswith(b'RIFF') and data[8:12] == b'WEBP'):
        return "WEBP signature mismatch"
    if ext == 'svg' and not data.lstrip().startswith(b'<svg'):
        return "SVG signature mismatch"
    if ext == 'pdf' and not data.startswith(b'%PDF'):
        return "PDF signature mismatch"
    if ext in {'zip', 'docx', 'xlsx', 'pptx'} and not data.startswith((b'PK\x03\x04', b'PK\x05\x06', b'PK\x07\x08')):
        return "ZIP document signature mismatch"
    if ext in {'doc', 'xls', 'ppt'} and not data.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
        return "Office document signature mismatch"
    if ext in {'mp4', 'mov'} and data[4:8] != b'ftyp':
        return "video signature mismatch"
    return ""


def _asset_bytes_valid(data, url):
    return not _asset_invalid_reason(data, url)


def _get_session():
    import requests
    return requests.Session()


_LOCAL_ASSET_CACHE: dict[str, Path] = {}


def prime_downloaded_asset_cache(jsonl_dirs):
    """Map scraped asset URLs to files already downloaded by scraper runs."""
    _LOCAL_ASSET_CACHE.clear()
    dirs = jsonl_dirs if isinstance(jsonl_dirs, list) else [jsonl_dirs]
    for jsonl_dir in dirs:
        root = Path(jsonl_dir)
        if not root.exists():
            continue
        scan_dirs = [root]
        if root.is_dir():
            scan_dirs.extend(p for p in root.iterdir() if p.is_dir() and not p.name.startswith("."))
        for scan_dir in scan_dirs:
            asset_file = scan_dir / "assets_unique.jsonl"
            if not asset_file.exists():
                continue
            try:
                lines = asset_file.read_text(encoding="utf-8", errors="ignore").splitlines()
            except OSError:
                continue
            for line in lines:
                if not line.strip():
                    continue
                try:
                    rec = json.loads(line)
                except json.JSONDecodeError:
                    continue
                local_path = rec.get("local_path")
                if not local_path:
                    continue
                path = scan_dir / str(local_path)
                if not path.is_file():
                    continue
                for key in (rec.get("url"), rec.get("download_url")):
                    if key:
                        _LOCAL_ASSET_CACHE[str(key)] = path


def _download_asset(url, headers=None, retries=1, delay=0.2):
    """Download an asset from a URL and return the file content.
    
    Retries on 429 with exponential backoff and jitter.
    """
    import time
    import random
    if not url:
        return None
    cached = _LOCAL_ASSET_CACHE.get(str(url))
    if cached and cached.is_file():
        data = cached.read_bytes()
        invalid_reason = _asset_invalid_reason(data, url)
        if not invalid_reason:
            return data
    request_headers = {'User-Agent': FETCH_UA}
    if headers:
        request_headers.update(headers)
    exc = None
    for attempt in range(1, retries + 1):
        try:
            session = _get_session()
            resp = session.get(url, headers=request_headers, timeout=30, allow_redirects=True)
            content_type = resp.headers.get('content-type', '').split(';', 1)[0].strip()
            if resp.status_code == 404:
                raise PermanentAssetDownloadError(
                    f"HTTP 404 ({content_type or 'unknown content-type'}, final={resp.url})"
                )
            if 400 <= resp.status_code < 500 and resp.status_code != 429:
                raise PermanentAssetDownloadError(
                    f"HTTP {resp.status_code} ({content_type or 'unknown content-type'}, final={resp.url})"
                )
            if resp.status_code == 429 and attempt < retries:
                wait = delay * (3 ** (attempt - 1)) + random.uniform(0, 2)
                if attempt < retries:
                    time.sleep(wait)
                continue
            resp.raise_for_status()
            data = resp.content
            invalid_reason = _asset_invalid_reason(data, url)
            if invalid_reason:
                raise PermanentAssetDownloadError(
                    f"{invalid_reason} ({content_type or 'unknown content-type'}, final={resp.url})"
                )
            return data
        except PermanentAssetDownloadError as e:
            log.warning(f"Unavailable asset {url}: {e}")
            return None
        except Exception as e:
            exc = e
            if attempt < retries:
                wait = delay * (3 ** (attempt - 1)) + random.uniform(0, 2)
                time.sleep(wait)
                continue
    if exc:
        log.error(f"Error downloading {url}: {exc}")
    return None


def _asset_storage_path(asset_dir, db_kind, filename):
    subdir = asset_dir / db_kind
    subdir.mkdir(parents=True, exist_ok=True)
    return subdir / filename, f"{db_kind}/{filename}"


def add_asset(conn, url, asset_type, entity_type, entity_id, role=None, sort_order=0):
    """Add or update an asset in the database and return its ID."""
    if not url:
        return None
    
    asset_dir = ASSETS_DIR
    asset_dir.mkdir(parents=True, exist_ok=True)
    
    # Try to download
    data = _download_asset(url)
    if not data:
        return None
    
    # Check for duplicates using checksum
    checksum = _sha256_bytes(data)
    existing = conn.execute(
        'SELECT id FROM assets WHERE checksum=?', 
        (checksum,)
    ).fetchone()
    
    if existing:
        asset_id = existing['id']
    else:
        # Save the file
        ext = _ext_from_url(url)
        filename = f"{checksum}.{ext}"
        # Determine real asset kind from file extension, not from entity type
        real_kind = _guess_kind_from_ext(ext)
        # Map to valid asset kind for CHECK constraint
        db_kind = real_kind if real_kind in ('image', 'document', 'pdf', 'video', 'other') else 'other'
        file_path, db_path = _asset_storage_path(asset_dir, db_kind, filename)
        if not file_path.exists():
            file_path.write_bytes(data)
        
        cur = conn.execute('''
            INSERT INTO assets (filename, original_filename, path, mime_type, size, kind, 
                                alt_text, caption, source_url, storage_status, is_external,
                                checksum, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        ''', (filename, filename, db_path, _guess_mime_from_ext(url), 
              len(data), db_kind, '', '', url[:2000], 'local', 0, checksum))
        asset_id = cur.lastrowid
    
    # Add asset link
    if asset_id:
        conn.execute('''
            INSERT OR REPLACE INTO asset_links 
            (entity_type, entity_id, asset_id, role, sort_order, created_at)
            VALUES (?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        ''', (entity_type, entity_id, asset_id, role, sort_order))
    
    return asset_id


def download_all_assets(conn, jsonl_dirs):
    """Download all unique assets with parallel workers + tqdm."""
    import threading
    
    if not isinstance(jsonl_dirs, list):
        jsonl_dirs = [jsonl_dirs]
    
    # Phase 1: Scan directories for assets_unique.jsonl
    seen_urls = set()
    url_list = []
    jsonl_files = []
    for jsonl_dir in jsonl_dirs:
        dirs_to_scan = [Path(jsonl_dir)]
        if os.path.isdir(jsonl_dir):
            for sub in os.listdir(jsonl_dir):
                sub_path = Path(jsonl_dir) / sub
                if sub_path.is_dir() and not sub.startswith('.'):
                    dirs_to_scan.append(sub_path)
        for scan_dir in dirs_to_scan:
            for fname in scan_dir.glob('assets_unique.jsonl'):
                jsonl_files.append(fname)
                break
    
    for fname in tqdm(jsonl_files, desc="Scanning JSONL", unit="file", leave=False):
        with open(fname, 'r', encoding='utf-8') as f:
            for line in f:
                if not line.strip():
                    continue
                rec = json.loads(line)
                url = rec.get('download_url') or rec.get('url')
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    url_list.append(url)
    
    if not url_list:
        return 0
    
    # Phase 2: Parallel download with per-thread + total progress
    lock = threading.Lock()
    total_pbar = tqdm(total=len(url_list), desc="Total assets", unit="file", position=0)
    
    def _worker(url, slot):
        """Worker with per-thread tqdm bar for streaming download."""
        try:
            data = _download_asset(url)
            with lock:
                total_pbar.update(1)
            return (url, data, None)
        except Exception as e:
            with lock:
                total_pbar.update(1)
            return (url, None, str(e))
    
    results = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=16) as pool:
        futures = {pool.submit(_worker, url, i % 8): url for i, url in enumerate(url_list)}
        for f in concurrent.futures.as_completed(futures):
            results.append(f.result())
    
    total_pbar.close()
    
    # Phase 3: Process results sequentially (DB ops)
    asset_dir = ASSETS_DIR
    asset_dir.mkdir(parents=True, exist_ok=True)
    
    count = 0
    skipped = 0
    for url, data, err in tqdm(results, desc="Storing", unit="file", leave=False):
        if not data:
            skipped += 1
            continue
        checksum = _sha256_bytes(data)
        existing = conn.execute(
            'SELECT id FROM assets WHERE checksum=?', (checksum,)
        ).fetchone()
        if existing:
            continue
        ext = _ext_from_url(url)

        # Extract title from document content if available
        title = _extract_title_from_document(data, ext)

        # Use title if available (first 30 chars), otherwise fallback to checksum
        if title:
            safe_title = _safe_stem(title[:30])
            filename = f"{safe_title or checksum[:12]}-{checksum[:12]}.{ext}"
        else:
            # Fallback to checksum if title extraction fails
            filename = f"{checksum}.{ext}"
        # Determine real asset kind from extension instead of hardcoding 'image'
        real_kind = _guess_kind_from_ext(ext)
        db_kind = real_kind if real_kind in ('image', 'document', 'pdf', 'video', 'other') else 'other'
        path, db_path = _asset_storage_path(asset_dir, db_kind, filename)
        if not path.exists():
            path.write_bytes(data)
        conn.execute('''
            INSERT INTO assets (filename, original_filename, path, mime_type, size, kind, 
                                alt_text, caption, source_url, storage_status, is_external,
                                checksum, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        ''', (filename, filename, db_path, _guess_mime_from_ext(url), len(data), db_kind,
              '', '', url[:2000], 'local', 0, checksum))
        count += 1
    
    if skipped:
        log.info(f"  Skipped {skipped} unavailable/invalid asset(s)")
    conn.commit()
    return count
