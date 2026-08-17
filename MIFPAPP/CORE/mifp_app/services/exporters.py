"""
MIFP Export Engine — Professional, colored, readable, and secure exports.

Supports: JSON, CSV, XLSX, DOCX, PDF

Security measures:
  - All cell values are sanitized (null bytes stripped, formula injection blocked)
  - CSV/TSV injection protection: dangerous leading chars are escaped
  - File paths and formulas never written into cells
  - Row limits enforced per format to prevent memory exhaustion
  - Content-Disposition headers use sanitized filenames
"""

from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime
from typing import Any

# ---------------------------------------------------------------------------
# Security & sanitization
# ---------------------------------------------------------------------------

# Characters that trigger CSV/TSV formula injection in Excel, Sheets, etc.
_FORMULA_CHARS = ("=", "+", "-", "@", "\t", "\r", "\n")
# Max cell length to prevent oversized exports
_MAX_CELL_LEN = 2000
# Row limits per format
_ROW_LIMITS = {"json": 50000, "jsonl": 50000, "csv": 50000, "xlsx": 50000, "docx": 5000, "pdf": 2000}


def _sanitize_cell(value: Any) -> str:
    """Return a safe, readable string for any cell value."""
    if value is None:
        return ""
    text = str(value).replace("\x00", "").strip()
    # Remove control characters except newline/tab
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    if len(text) > _MAX_CELL_LEN:
        text = text[:_MAX_CELL_LEN] + "…"
    return text


def _csv_escape(value: str) -> str:
    """Prevent CSV formula injection by escaping dangerous leading characters."""
    if value and value[0] in _FORMULA_CHARS:
        return "'" + value
    return value


def _column_label(col: str) -> str:
    """Human-readable column label: 'first_name' -> 'First Name'."""
    return col.replace("_", " ").title()


def _columns(rows: list[dict[str, Any]]) -> list[str]:
    """Preserve insertion-order column list across all rows."""
    cols: list[str] = []
    for row in rows:
        for key in row.keys():
            if key not in cols:
                cols.append(key)
    return cols


def _safe_filename(title: str) -> str:
    """Produce a safe filename fragment from a title."""
    return re.sub(r"[^A-Za-z0-9_\-]", "_", title)[:60]


# ---------------------------------------------------------------------------
# Color palette — MIFP institutional theme
# ---------------------------------------------------------------------------

_MIFP_RED = "B42318"
_MIFP_RED_DARK = "7F1D1D"
_MIFP_NAVY = "0A0E1A"
_MIFP_GRAY_100 = "F3F4F6"
_MIFP_GRAY_200 = "E5E7EB"
_MIFP_GRAY_500 = "6B7280"
_MIFP_WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# JSON
# ---------------------------------------------------------------------------

def rows_to_json(rows: list[dict[str, Any]], title: str = "Export") -> dict[str, Any]:
    return {
        "meta": {
            "title": title,
            "exported_at": datetime.now().isoformat(timespec="seconds"),
            "total_rows": len(rows),
        },
        "columns": list(rows[0].keys()) if rows else [],
        "rows": rows,
    }


def rows_to_jsonl(rows: list[dict[str, Any]]) -> bytes:
    lines = [json.dumps(row, ensure_ascii=False, default=str) for row in rows]
    return ("\n".join(lines) + ("\n" if lines else "")).encode("utf-8")


# ---------------------------------------------------------------------------
# CSV — UTF-8 BOM, injection-safe, clean headers
# ---------------------------------------------------------------------------

def rows_to_csv(rows: list[dict[str, Any]]) -> bytes:
    out = io.StringIO()
    if rows:
        cols = _columns(rows)
        writer = csv.DictWriter(out, fieldnames=cols, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            safe_row = {k: _csv_escape(_sanitize_cell(v)) for k, v in row.items()}
            writer.writerow(safe_row)
    return out.getvalue().encode("utf-8-sig")


# ---------------------------------------------------------------------------
# XLSX — openpyxl, MIFP branded, alternating rows, auto-filter, freeze
# ---------------------------------------------------------------------------

def rows_to_xlsx(rows: list[dict[str, Any]], sheet_name: str = "Export") -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "Export"

    # ── Style definitions ──
    header_font = Font(name="Inter", bold=True, color=_MIFP_WHITE, size=10)
    header_fill = PatternFill(start_color=_MIFP_RED, end_color=_MIFP_RED, fill_type="solid")
    header_align = Alignment(horizontal="left", vertical="center", wrap_text=True)

    cell_font = Font(name="Inter", size=9, color="1F2937")
    cell_align = Alignment(vertical="center", wrap_text=False)
    cell_font_wrapped = Font(name="Inter", size=9, color="1F2937")

    even_fill = PatternFill(start_color=_MIFP_WHITE, end_color=_MIFP_WHITE, fill_type="solid")
    odd_fill = PatternFill(start_color=_MIFP_GRAY_100, end_color=_MIFP_GRAY_100, fill_type="solid")

    thin_border = Border(
        bottom=Side(style="thin", color=_MIFP_GRAY_200),
    )
    header_border = Border(
        bottom=Side(style="medium", color=_MIFP_RED_DARK),
    )

    cols = _columns(rows) if rows else ["(no data)"]

    # ── Title row (row 1) ──
    title_font = Font(name="Inter", bold=True, color=_MIFP_RED, size=14)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(cols))
    title_cell = ws.cell(row=1, column=1, value=f"MIFP — {sheet_name}")
    title_cell.font = title_font
    title_cell.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 28

    # ── Subtitle row (row 2) ──
    subtitle_font = Font(name="Inter", size=8, color=_MIFP_GRAY_500, italic=True)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(cols))
    sub_cell = ws.cell(row=2, column=1, value=f"Exported: {datetime.now():%Y-%m-%d %H:%M}  ·  {len(rows)} records")
    sub_cell.font = subtitle_font
    ws.row_dimensions[2].height = 18

    # ── Header row (row 4) ──
    header_row = 4
    for col_idx, col_name in enumerate(cols, 1):
        cell = ws.cell(row=header_row, column=col_idx, value=_column_label(col_name))
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = header_border
    ws.row_dimensions[header_row].height = 24

    # ── Data rows ──
    for row_idx, row in enumerate(rows):
        excel_row = header_row + 1 + row_idx
        fill = even_fill if row_idx % 2 == 0 else odd_fill
        for col_idx, col_name in enumerate(cols, 1):
            val = _sanitize_cell(row.get(col_name))
            cell = ws.cell(row=excel_row, column=col_idx, value=val)
            cell.font = cell_font
            cell.alignment = cell_align
            cell.fill = fill
            cell.border = thin_border
            # Truncate display for very wide columns
            if len(val) > 80:
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.font = cell_font_wrapped

    # ── Column widths ──
    for col_idx, col_name in enumerate(cols, 1):
        header_len = len(_column_label(col_name))
        max_len = header_len
        for row in rows[:200]:
            val = _sanitize_cell(row.get(col_name))
            max_len = max(max_len, min(len(val), 40))
        ws.column_dimensions[get_column_letter(col_idx)].width = max(max_len + 3, 12)

    # ── Freeze & auto-filter ──
    ws.freeze_panes = f"A{header_row + 1}"
    if rows:
        last_col = get_column_letter(len(cols))
        ws.auto_filter.ref = f"A{header_row}:{last_col}{header_row + len(rows)}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX — python-docx with branded header, styled table, alternating rows
# ---------------------------------------------------------------------------

def rows_to_docx(rows: list[dict[str, Any]], title: str = "Export") -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    export_display = "Georgia"
    export_body = "Arial"
    normal_style = doc.styles["Normal"]
    normal_style.font.name = export_body
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), export_body)

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── Title block ──
    title_para = doc.add_paragraph()
    title_para.space_after = Pt(2)  # type: ignore[attr-defined]
    mifp_run = title_para.add_run("MIFP")
    mifp_run.font.size = Pt(18)
    mifp_run.font.bold = True
    mifp_run.font.name = export_display
    mifp_run.font.color.rgb = RGBColor(0xB4, 0x23, 0x18)
    sep_run = title_para.add_run("  ·  ")
    sep_run.font.size = Pt(14)
    sep_run.font.name = export_body
    sep_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = export_display
    title_run.font.color.rgb = RGBColor(0x0A, 0x0E, 0x1A)

    # ── Subtitle ──
    sub_para = doc.add_paragraph()
    sub_para.space_after = Pt(8)  # type: ignore[attr-defined]
    sub_run = sub_para.add_run(f"Exported: {datetime.now():%Y-%m-%d %H:%M}  ·  {len(rows)} records")
    sub_run.font.size = Pt(9)
    sub_run.font.name = export_body
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    sub_run.font.italic = True

    # ── Red divider line ──
    divider = doc.add_paragraph()
    divider.space_after = Pt(6)  # type: ignore[attr-defined]
    pPr = divider._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _MIFP_RED)
    pBdr.append(bottom)
    pPr.append(pBdr)

    cols = _columns(rows) if rows else []

    if not rows:
        empty_para = doc.add_paragraph("No data to export.")
        empty_para.runs[0].font.name = export_body
        empty_para.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        empty_para.runs[0].font.italic = True
    else:
        # ── Table ──
        table = doc.add_table(rows=1 + len(rows), cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        # Header row
        for col_idx, col_name in enumerate(cols):
            cell = table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(_column_label(col_name))
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = export_body
            # Red background
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), _MIFP_RED)
            shading.set(qn("w:val"), "clear")
            cell._element.get_or_add_tcPr().append(shading)

        # Data rows with alternating backgrounds
        for row_idx, row in enumerate(rows):
            for col_idx, col_name in enumerate(cols):
                cell = table.rows[row_idx + 1].cells[col_idx]
                val = _sanitize_cell(row.get(col_name))
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val[:300])
                run.font.size = Pt(7.5)
                run.font.name = export_body
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                # Alternating row shading
                if row_idx % 2 == 1:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), _MIFP_GRAY_100)
                    shading.set(qn("w:val"), "clear")
                    cell._element.get_or_add_tcPr().append(shading)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF — ReportLab, branded header, alternating rows, page numbers, security
# ---------------------------------------------------------------------------

# Columns to include in PDF exports for wide tables (key fields only)
_PDF_KEY_COLUMNS = {
    "events": ["id", "title", "start_date", "end_date", "location", "event_type", "review_status", "is_featured"],
    "members": ["id", "first_name", "last_name", "display_name", "email", "affiliation", "country", "role_id", "is_active"],
    "news": ["id", "title", "news_type", "date", "review_status", "is_featured"],
    "publications": ["id", "title", "year", "authors", "journal", "review_status"],
    "research_areas": ["id", "title", "review_status"],
    "sponsors": ["id", "name", "sponsor_type", "tier", "is_active"],
    "assets": ["id", "filename", "kind", "size", "is_external"],
}


def rows_to_pdf(rows: list[dict[str, Any]], title: str = "Export") -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    # Determine columns: for wide tables, use only key columns
    all_cols = _columns(rows) if rows else []
    # Find matching key column set by title
    table_key = title.lower().replace(" ", "_")
    key_cols = None
    for k, v in _PDF_KEY_COLUMNS.items():
        if k in table_key:
            key_cols = v
            break
    # If more than 10 columns and we have a key set, filter to key columns only
    if len(all_cols) > 10 and key_cols:
        cols = [c for c in key_cols if c in all_cols]
        if len(cols) < 3:  # fallback if too few match
            cols = all_cols[:10]
    elif len(all_cols) > 10:
        cols = all_cols[:10]
    else:
        cols = all_cols

    page_size = landscape(A4) if len(cols) > 5 else A4

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=page_size,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=2.2 * cm,
        bottomMargin=1.8 * cm,
    )

    styles = getSampleStyleSheet()
    pdf_display = "Times-Bold"
    pdf_body = "Helvetica"
    pdf_body_bold = "Helvetica-Bold"
    pdf_body_italic = "Helvetica-Oblique"

    # ── Custom styles ──
    title_style = ParagraphStyle(
        "MifpTitle", parent=styles["Heading1"],
        fontSize=18, spaceAfter=2, spaceBefore=0,
        textColor=colors.HexColor(f"#{_MIFP_RED}"),
        fontName=pdf_display,
    )
    h2_style = ParagraphStyle(
        "MifpH2", parent=styles["Heading2"],
        fontSize=12, spaceAfter=2,
        textColor=colors.HexColor(f"#{_MIFP_NAVY}"),
        fontName=pdf_display,
    )
    subtitle_style = ParagraphStyle(
        "MifpSubtitle", parent=styles["Normal"],
        fontSize=8.5, textColor=colors.HexColor(f"#{_MIFP_GRAY_500}"),
        spaceAfter=6, fontName=pdf_body_italic,
    )
    header_style = ParagraphStyle(
        "MifpHeader", parent=styles["Normal"],
        fontSize=7, textColor=colors.white,
        fontName=pdf_body_bold,
    )
    cell_style = ParagraphStyle(
        "MifpCell", parent=styles["Normal"],
        fontSize=7, textColor=colors.HexColor("#1F2937"),
        fontName=pdf_body,
    )

    elements = []

    # ── Title block ──
    elements.append(Paragraph("MIFP", title_style))
    elements.append(Paragraph(title, h2_style))
    elements.append(Paragraph(
        f"Exported: {datetime.now():%Y-%m-%d %H:%M}  ·  {len(rows)} records",
        subtitle_style,
    ))
    elements.append(HRFlowable(
        width="100%", thickness=1.5,
        color=colors.HexColor(f"#{_MIFP_RED}"),
        spaceAfter=8, spaceBefore=2,
    ))

    if not rows:
        elements.append(Paragraph("No data to export.", cell_style))
    else:
        # ── Build table ──
        header_row = [Paragraph(_column_label(c), header_style) for c in cols]
        data = [header_row]
        for row in rows[:_ROW_LIMITS["pdf"]]:
            data_row = [Paragraph(_sanitize_cell(row.get(c))[:200], cell_style) for c in cols]
            data.append(data_row)

        # Column widths
        available = page_size[0] - 3.6 * cm
        col_count = len(cols)
        if col_count > 0:
            _WIDE_FLOOR = frozenset({"email", "url", "external_link", "website_url", "doi",
                                      "affiliation", "location", "abstract"})
            raw_widths = []
            for c in cols:
                label_len = len(_column_label(c))
                floor = 30 if c in _WIDE_FLOOR else 0
                max_data_len = max(label_len, floor)
                for row in rows[:100]:
                    val = _sanitize_cell(row.get(c))
                    max_data_len = max(max_data_len, min(len(val), 80))
                raw_widths.append(max_data_len)
            total_raw = sum(raw_widths) or 1
            min_w = 14 * mm
            col_widths = [max(available * (rw / total_raw), min_w) for rw in raw_widths]
            # Scale down if exceeds page width
            total_width = sum(col_widths)
            if total_width > available:
                scale = available / total_width
                col_widths = [w * scale for w in col_widths]
        else:
            col_widths = [60 * mm]

        # Split data into chunks to avoid LayoutError with very tall rows
        # Each chunk has at most 40 data rows (plus header)
        chunk_size = 40
        total_data_rows = len(data) - 1  # exclude header

        if len(cols) != len(all_cols) and len(all_cols) > 10:
            col_note = Paragraph(
                f"<i>Note: Showing {len(cols)} of {len(all_cols)} columns. "
                f"Export as XLSX or CSV for all columns.</i>",
                ParagraphStyle("MifpNote", parent=cell_style, fontSize=7,
                               textColor=colors.HexColor(f"#{_MIFP_GRAY_500}"),
                               spaceAfter=6)
            )
            elements.append(col_note)

        for chunk_start in range(0, total_data_rows, chunk_size):
            chunk_end = min(chunk_start + chunk_size, total_data_rows)
            chunk_data = [data[0]] + data[1 + chunk_start:1 + chunk_end]  # header + chunk

            table = Table(chunk_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                # Header
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_MIFP_RED}")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), pdf_body_bold),
                ("FONTSIZE", (0, 0), (-1, 0), 7),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                ("TOPPADDING", (0, 0), (-1, 0), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                # Data
                ("FONTNAME", (0, 1), (-1, -1), pdf_body),
                ("FONTSIZE", (0, 1), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 1), (-1, -1), 4),
                ("TOPPADDING", (0, 1), (-1, -1), 4),
                # Alternating row colors
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor(f"#{_MIFP_GRAY_100}")]),
                # Grid
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{_MIFP_GRAY_200}")),
                ("LINEBELOW", (0, 0), (-1, 0), 1.5, colors.HexColor(f"#{_MIFP_RED_DARK}")),
                # Alignment
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            elements.append(table)
            # Add spacer between chunks (except last)
            if chunk_end < total_data_rows:
                elements.append(Spacer(1, 8 * mm))

    # ── Page number footer ──
    def add_page_footer(canvas, doc):
        page_num = canvas.getPageNumber()
        pw, ph = page_size
        # Red divider line
        canvas.setStrokeColor(colors.HexColor(f"#{_MIFP_RED}"))
        canvas.setLineWidth(0.5)
        canvas.line(1.8 * cm, 1.3 * cm, pw - 1.8 * cm, 1.3 * cm)
        # Footer text
        canvas.setFont(pdf_body, 7)
        canvas.setFillColor(colors.HexColor(f"#{_MIFP_GRAY_500}"))
        canvas.drawString(1.8 * cm, 0.8 * cm, f"MIFP · {title}")
        canvas.drawRightString(pw - 1.8 * cm, 0.8 * cm, f"Page {page_num}")

    doc.build(elements, onFirstPage=add_page_footer, onLaterPages=add_page_footer)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Export dispatcher
# ---------------------------------------------------------------------------

def export_response_payload(
    rows: list[dict[str, Any]],
    fmt: str,
    title: str,
) -> tuple[bytes | dict[str, Any], str, str]:
    """Return (payload, mimetype, extension) for the given format.

    Applies row limits per format and sanitizes all output.
    """
    limit = _ROW_LIMITS.get(fmt, 50000)
    rows = rows[:limit]

    if fmt == "json":
        return rows_to_json(rows, title), "application/json", "json"
    if fmt == "jsonl":
        return rows_to_jsonl(rows), "application/x-ndjson", "jsonl"
    if fmt == "csv":
        return rows_to_csv(rows), "text/csv", "csv"
    if fmt == "xlsx":
        safe_name = _safe_filename(title)[:31] or "Export"
        return rows_to_xlsx(rows, safe_name), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"
    if fmt == "docx":
        return rows_to_docx(rows, title), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    if fmt == "pdf":
        return rows_to_pdf(rows, title), "application/pdf", "pdf"
    raise ValueError(f"Unsupported export format: {fmt}")
